import re
from dataclasses import dataclass
from io import BytesIO

from docx import Document


@dataclass
class DocumentContent:
    paragraphs: list[str]
    tables: list[list[list[str]]]
    text: str


class DocumentService:
    @staticmethod
    def CleanParagraph(text: str) -> str | None:
        text = text.strip()

        if not text:
            return None

        # Page markers produced by the source document.
        if re.fullmatch(r"\d+\s*#", text):
            return None

        return text

    @staticmethod
    def ExtractDOCX(content: bytes) -> DocumentContent:
        document = Document(BytesIO(content))

        paragraphs = []

        for paragraph in document.paragraphs:
            text = DocumentService.CleanParagraph(paragraph.text)
            if text is not None:
                paragraphs.append(text)

        tables: list[list[list[str]]] = []
        for table in document.tables:
            parsed_table: list[list[str]] = []
            for row in table.rows:
                parsed_row = [cell.text.strip() for cell in row.cells]
                parsed_table.append(parsed_row)
            tables.append(parsed_table)
        text_parts = paragraphs.copy()
        for table in tables:
            for row in table:
                text_parts.extend(cell for cell in row if cell)

        return DocumentContent(
            paragraphs=paragraphs,
            tables=tables,
            text="\n".join(text_parts),
        )
